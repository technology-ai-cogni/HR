def extract_text_from_file(file_obj):
    # Determine file extension
    file_name = file_obj.name if hasattr(file_obj, "name") else file_obj
    if file_name.lower().endswith('.pdf'):
        from PyPDF2 import PdfReader
        reader = PdfReader(file_obj)
        text = ""
        for page in reader.pages:
            text += (page.extract_text() or "") + "\n"
        return text
    elif file_name.lower().endswith('.docx'):
        import docx
        doc = docx.Document(file_obj)
        full_text = []

        # 1. Extract headers from sections (often contains contact info)
        for section in doc.sections:
            if section.header:
                for p in section.header.paragraphs:
                    if p.text.strip():
                        full_text.append(p.text.strip())
                for t in section.header.tables:
                    for row in t.rows:
                        for cell in row.cells:
                            if cell.text.strip():
                                full_text.append(cell.text.strip())

        # 2. Extract body paragraphs
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                full_text.append(paragraph.text)

        # 3. Extract body tables (often contains structured resume info)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        full_text.append(cell.text.strip())

        # 4. Extract footers from sections
        for section in doc.sections:
            if section.footer:
                for p in section.footer.paragraphs:
                    if p.text.strip():
                        full_text.append(p.text.strip())

        return "\n".join(full_text)
    else:
        raise ValueError("Unsupported file type")
